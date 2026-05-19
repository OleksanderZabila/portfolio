"""Generate a professional CV for Aleksandr Zabila based on DOU.ua standards."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Calibri"
BLACK = RGBColor(0, 0, 0)
DARK = RGBColor(0x33, 0x33, 0x33)


def set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    style.font.color.rgb = DARK
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)


def set_margins(doc, cm=1.6):
    for s in doc.sections:
        s.top_margin = Cm(cm)
        s.bottom_margin = Cm(cm)
        s.left_margin = Cm(cm)
        s.right_margin = Cm(cm)


def add_spacer(doc, pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.add_run("").font.size = Pt(pt)


def add_name(doc, name):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = BLACK


def add_title(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.font.size = Pt(12)
    r.font.color.rgb = DARK


def add_contacts(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.font.size = Pt(10)
        r.font.color.rgb = DARK


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = BLACK
    # add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_body(doc, text, size=11, bold=False, italic=False, after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = DARK
    return p


def add_job_header(doc, role, company, dates):
    """One line: 'Role — Company [tab] Dates'"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    # tab stop on right
    from docx.enum.text import WD_TAB_ALIGNMENT
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(17.5), WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run(role + " — ")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BLACK
    rc = p.add_run(company)
    rc.bold = True
    rc.font.size = Pt(11)
    rc.font.color.rgb = BLACK
    rd = p.add_run("\t" + dates)
    rd.italic = True
    rd.font.size = Pt(10)
    rd.font.color.rgb = DARK


def add_tech_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = DARK


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(item)
        r.font.size = Pt(11)
        r.font.color.rgb = DARK


def add_skill_row(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    rl = p.add_run(label + "  ")
    rl.bold = True
    rl.font.size = Pt(11)
    rl.font.color.rgb = BLACK
    rv = p.add_run(value)
    rv.font.size = Pt(11)
    rv.font.color.rgb = DARK


def build():
    doc = Document()
    set_default_font(doc)
    set_margins(doc, 1.6)

    # Header
    add_name(doc, "Aleksandr Zabila")
    add_title(doc, "Junior Python Developer")
    add_contacts(doc, [
        "+380 50 838 4092  |  ukzabila@gmail.com  |  Cherkasy, Ukraine",
        "github.com/OleksanderZabila  |  t.me/ZabilaOleksandr",
        "Remote / Office  |  Open to relocate",
    ])

    # Summary
    add_section_heading(doc, "Summary")
    add_body(doc,
        "Computer Engineering Master's student at MAUP (expected 2027) focused on "
        "Python backend development and PostgreSQL design. Built and deployed a "
        "production Telegram bot currently serving 2,100+ daily users, plus a "
        "desktop inventory system as a diploma project for a real client. Looking "
        "for a Junior Python Developer role where I can contribute to real products "
        "and grow within a strong engineering team.",
        after=2,
    )

    # Skills
    add_section_heading(doc, "Skills")
    add_skill_row(doc, "Languages:", "Python (primary), C++, SQL, JavaScript (basic)")
    add_skill_row(doc, "Backend:", "Django (ORM, admin, management commands), python-telegram-bot, asyncio, REST API integration (GitHub, Telegram)")
    add_skill_row(doc, "Databases:", "PostgreSQL (schema design, queries, normalisation), SQLite")
    add_skill_row(doc, "Frontend:", "HTML, CSS, Tailwind CSS, vanilla JavaScript, responsive design")
    add_skill_row(doc, "Tools:", "Git, GitHub, PyCharm, VS Code, pgAdmin, Tkinter, matplotlib, ngrok")
    add_skill_row(doc, "Concepts:", "OOP, CRUD, role-based access control, caching, environment-based config, deployment basics")
    add_skill_row(doc, "Spoken:", "Ukrainian — Native  |  English — A2 (improving)")

    # Projects
    add_section_heading(doc, "Projects")

    # 1. Blackout Bot
    add_job_header(doc, "Blackout Schedule Bot", "Production Telegram Bot", "2025 — Present")
    add_tech_line(doc, "Python · python-telegram-bot · matplotlib · SQLite · asyncio")
    add_bullets(doc, [
        "Designed and deployed a Telegram bot serving 2,100+ daily active users with personalised Ukrainian power-outage schedules.",
        "Generated live hour-by-hour timeline charts with matplotlib showing the current time position and outage windows per energy group.",
        "Implemented admin analytics: new-user notifications, daily-growth tracking, total user counter, real-time chat-id capture.",
        "Maintained 99%+ uptime since launch; source private, currently in production.",
    ])

    # 2. Portfolio Website
    add_job_header(doc, "Personal Portfolio Website", "Solo Project", "May 2026")
    add_tech_line(doc, "Django · Tailwind CSS · GitHub API · Telegram Bot API · WhiteNoise")
    add_bullets(doc, [
        "Built a single-page portfolio site from scratch with 6 Django models and a fully customised admin panel.",
        "Integrated the GitHub REST API with 1-hour file-based caching to auto-sync 19+ repositories and language statistics.",
        "Wired Telegram notifications for contact-form submissions using the Bot API with HTML-formatted messages.",
        "Implemented dark-neon UI with animated gradients, scroll-reveal, typewriter effect, and project image galleries.",
        "Source: github.com/OleksanderZabila/portfolio",
    ])

    # 3. Auto Pidkliuch
    add_job_header(doc, "Auto Pidkliuch — Inventory Management System", "Diploma Project", "Feb 2025 — May 2025")
    add_tech_line(doc, "Python · PostgreSQL · Tkinter · psycopg2")
    add_bullets(doc, [
        "Designed an 8-table normalised PostgreSQL schema for an auto-parts store covering products, suppliers, customers, and sales.",
        "Implemented full CRUD across 4 core entities and a role-based access control system (Admin / Cashier).",
        "Built a Tkinter analytics dashboard for daily sales tracking and reporting; reduced inventory lookup time from manual to under 1 second.",
        "Owned the project end-to-end from schema design through GUI to handover.",
    ])

    # 4. Weather Bot
    add_job_header(doc, "Weather Forecast Telegram Bot", "Hobby Project", "Dec 2023")
    add_tech_line(doc, "Python · python-telegram-bot · matplotlib · OpenWeather API")
    add_bullets(doc, [
        "Built a Telegram bot delivering daily forecasts with matplotlib-rendered temperature charts.",
        "Implemented command parsing, per-user city storage, and graceful API error handling.",
        "Source: github.com/OleksanderZabila/weather_telegram_bot-master",
    ])

    # 5. Transcriber Bot
    add_job_header(doc, "Transcriber Bot", "Interview Challenge", "Oct 2025")
    add_tech_line(doc, "Python · python-telegram-bot · speech-to-text")
    add_bullets(doc, [
        "Telegram bot accepting voice/audio messages and returning transcribed text, built as a take-home interview task.",
        "Source: github.com/OleksanderZabila/transcriber-bot",
    ])

    # Education
    add_section_heading(doc, "Education")
    add_job_header(doc, "Master's in Computer Engineering", "Interregional Academy of Personnel Management (MAUP)", "2021 — 2027 (expected)")
    add_tech_line(doc, "Specialty 121 · Online · Bachelor's completed within the same program · Cherkasy, Ukraine")

    add_job_header(doc, "English Language Courses", "Online", "2024 — Present")
    add_tech_line(doc, "Current level: A2, actively improving for international team collaboration.")

    out = r"C:\Users\ukzab\Desktop\CV_Junior_Python_Developer_Aleksandr_Zabila_May2026_Eng.docx"
    doc.save(out)
    print("Saved:", out)


if __name__ == "__main__":
    build()
